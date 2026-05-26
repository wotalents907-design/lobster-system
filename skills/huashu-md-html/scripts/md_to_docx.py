#!/usr/bin/env python3
"""
md_to_docx.py — 把 markdown 加工成出版社级 docx

能力 4：md → 精美 docx（出版社审校 / 投稿 / 纸质书定稿）

用法
----
# 单文件
python3 md_to_docx.py article.md
python3 md_to_docx.py article.md -o article.docx
python3 md_to_docx.py article.md --images-dir ./images

# 多文件合并成书（按文件名顺序）
python3 md_to_docx.py ch01.md ch02.md ch03.md -o book.docx

# 完整书模式（带封面 + 目录 + 页眉页脚）
python3 md_to_docx.py ch*.md --book \\
    --title "图解 Agent Skills" \\
    --subtitle "让 AI 记住你的工作方式" \\
    --author "花叔" \\
    -o book.docx

特性
----
- 中文友好的字体（思源宋体/思源黑体，回退 PingFang SC / Songti SC）
- 自动嵌入图片（识别 ![]() 内联和 ![][ref] 引用两种语法）
- 引用块按类型自动配色（💡 重点 / ✅ 建议 / ⚠️ 注意 / 普通）
- 代码块带左侧色边 + 浅灰底
- 表格干净边框 + 表头底色
- 章节首页专业版面（章号 + 标题 + 英文副标题 + 分隔线）
- 页眉书名 + 页脚自动页码
- 默认大 32 开（176×240 mm），适合纸质书

依赖
----
python3 -m pip install python-docx Pillow
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from PIL import Image as PILImage
except ImportError:
    print("缺少依赖：python3 -m pip install Pillow", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("缺少依赖：python3 -m pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 颜色 & 字体
# ============================================================
C_ORANGE = RGBColor(0xC2, 0x41, 0x0C)
C_TEAL = RGBColor(0x0E, 0x7C, 0x66)
C_INK = RGBColor(0x1A, 0x1A, 0x1A)
C_MUTED = RGBColor(0x6B, 0x6B, 0x6B)
C_ROSE = RGBColor(0xBE, 0x12, 0x3C)

FONT_CN_BODY = "思源宋体 CN"
FONT_CN_HEAD = "思源黑体 CN"
FONT_MONO = "JetBrains Mono"
FONT_EN = "Georgia"


# ============================================================
# 底层 OOXML 工具
# ============================================================
def set_run_font(run, size_pt=11, bold=False, italic=False, color=None,
                 font_cn=FONT_CN_BODY, font_en=None):
    run.font.name = font_en or font_cn
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    if font_en:
        rFonts.set(qn("w:ascii"), font_en)
        rFonts.set(qn("w:hAnsi"), font_en)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_horizontal_line(doc, color_hex="D6D0C4"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# 内联文本解析
# ============================================================
INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*)"           # **bold**
    r"|(\*([^*\n]+)\*)"            # *italic*
    r"|(`([^`]+)`)"                # `code`
    r"|(\[([^\]]+)\]\(([^)]+)\))"  # [text](url)
)


def add_inline_runs(paragraph, text, base_size=11):
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, size_pt=base_size)
        if m.group(2):  # bold
            r = paragraph.add_run(m.group(2))
            set_run_font(r, size_pt=base_size, bold=True)
        elif m.group(4):  # italic
            r = paragraph.add_run(m.group(4))
            set_run_font(r, size_pt=base_size, italic=True, color=C_MUTED)
        elif m.group(6):  # inline code
            r = paragraph.add_run(m.group(6))
            set_run_font(r, size_pt=base_size - 0.5, font_cn=FONT_MONO,
                         font_en=FONT_MONO, color=C_ORANGE)
        elif m.group(7):  # link
            r = paragraph.add_run(m.group(8))
            set_run_font(r, size_pt=base_size, color=C_ORANGE)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, size_pt=base_size)


# ============================================================
# Markdown 块解析
# ============================================================
def parse_blocks(md_text, image_refs):
    """
    把 md 文本切分成块。
    image_refs 字典：{"fig-1-1": "path/to/img.png", ...}（从引用式定义中预先提取）
    """
    # 去掉文件末尾的引用定义行（已经在 image_refs 里）
    md_text = re.sub(r"^\[[^\]]+\]:\s*\S+.*$", "", md_text, flags=re.MULTILINE)

    blocks = []
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            blocks.append({"kind": "hr"})
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            blocks.append({"kind": "heading", "level": len(m.group(1)),
                           "text": m.group(2).strip()})
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip() or "text"
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append({"kind": "code", "lang": lang, "lines": code_lines})
            continue

        # 引用式图片：> ![alt][ref]
        m_img = re.match(r"^>\s*!\[([^\]]+)\]\[([^\]]+)\]\s*$", stripped)
        if m_img:
            blocks.append({"kind": "image", "alt": m_img.group(1),
                           "ref": m_img.group(2), "via": "ref"})
            i += 1
            continue

        # 内联式图片（独立成行）：![alt](path)
        m_img2 = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m_img2:
            blocks.append({"kind": "image", "alt": m_img2.group(1),
                           "path": m_img2.group(2), "via": "inline"})
            i += 1
            continue

        # 引用块（含 callout/tip/warning）
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and (lines[i].strip().startswith(">") or
                             lines[i].strip() == ""):
                if lines[i].strip().startswith(">"):
                    quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                else:
                    quote_lines.append("")
                i += 1
            while quote_lines and not quote_lines[-1].strip():
                quote_lines.pop()
            blocks.append({"kind": "quote", "lines": quote_lines})
            continue

        # 表格
        if "|" in line and i + 1 < n and re.match(
                r"^\s*\|?\s*[-:]+", lines[i + 1].strip()):
            table_lines = [line]
            i += 2  # 跳过 separator
            while i < n and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            blocks.append({"kind": "table", "rows": table_lines})
            continue

        # 无序列表
        if re.match(r"^\s*-\s+", line):
            items = []
            while i < n and re.match(r"^\s*-\s+", lines[i]):
                items.append(re.sub(r"^\s*-\s+", "", lines[i]))
                i += 1
            blocks.append({"kind": "ul", "items": items})
            continue

        # 有序列表
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < n and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"kind": "ol", "items": items})
            continue

        # 普通段落（合并连续非空行）
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,6}\s|```|>|\|.*\||\s*-\s+|\s*\d+\.\s+|---\s*$|!\[)",
                lines[i].strip()):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"kind": "p",
                       "text": " ".join(l.strip() for l in para_lines)})

    return blocks


def extract_image_refs(md_text):
    """提取引用式图片定义：[fig-1-1]: path/to/img.png "alt" """
    refs = {}
    for m in re.finditer(
            r'^\[([^\]]+)\]:\s*(\S+)(?:\s+"([^"]+)")?',
            md_text, re.MULTILINE):
        refs[m.group(1)] = m.group(2)
    return refs


# ============================================================
# docx 渲染
# ============================================================
def add_heading(doc, level, text, chapter_label=None):
    if level == 1:
        # 章号小标
        if chapter_label:
            p = doc.add_paragraph()
            r = p.add_run(chapter_label)
            set_run_font(r, size_pt=11, color=C_MUTED, font_cn=FONT_CN_HEAD)
            p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        clean = re.sub(r"^(第\s*\d+\s*章|附录|后记)\s*", "", text)
        clean = re.sub(r"^\d{2}\s+", "", clean)
        r = p.add_run(clean if clean else text)
        set_run_font(r, size_pt=24, bold=True, color=C_INK, font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_after = Pt(14)
        # 橙色底部分隔线
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "8")
        bottom.set(qn("w:color"), "C2410C")
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif level == 2:
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size_pt=17, bold=True, color=C_INK,
                     font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)

    elif level == 3:
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size_pt=13.5, bold=True, color=C_ORANGE,
                     font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    else:
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size_pt=12, bold=True, color=C_INK,
                     font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_before = Pt(8)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    add_inline_runs(p, text, base_size=11)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(6)


def add_italic_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size_pt=11, italic=True, color=C_MUTED,
                 font_cn=FONT_CN_HEAD, font_en="Georgia")
    p.paragraph_format.space_after = Pt(14)


def add_image(doc, image_path, alt_text, max_width_inches=5.8):
    if not image_path or not Path(image_path).exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[未找到图片：{image_path}]")
        set_run_font(r, size_pt=9, italic=True, color=C_ROSE)
        return

    try:
        with PILImage.open(image_path) as img:
            iw, ih = img.size
    except Exception:
        iw, ih = 1400, 900

    width_inches = min(max_width_inches, iw / 150)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(width_inches))

    if alt_text:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(alt_text)
        set_run_font(r_cap, size_pt=9, italic=True, color=C_MUTED,
                     font_cn=FONT_CN_HEAD)
        p_cap.paragraph_format.space_after = Pt(14)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F5F5F0")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), "C2410C")
    pBdr.append(left)
    pPr.append(pBdr)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4

    r = p.add_run("\n".join(code_lines))
    set_run_font(r, size_pt=9.5, color=C_INK, font_cn=FONT_MONO,
                 font_en=FONT_MONO)


def add_quote_block(doc, lines):
    first_text = "\n".join(lines)
    if "⚠️" in first_text[:10] or "**⚠️" in first_text:
        border, fill = "BE123C", "FFE4E6"
    elif "💡" in first_text[:10] or "**💡" in first_text:
        border, fill = "D97706", "FEF3C7"
    elif "✅" in first_text[:10] or "**✅" in first_text:
        border, fill = "0E7C66", "CCFBF1"
    else:
        border, fill = "C2410C", "FFF7ED"

    for idx, line in enumerate(lines):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "20")
        left.set(qn("w:space"), "10")
        left.set(qn("w:color"), border)
        pBdr.append(left)
        pPr.append(pBdr)

        if idx == 0:
            p.paragraph_format.space_before = Pt(8)
        if idx == len(lines) - 1:
            p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        add_inline_runs(p, line, base_size=10.5)


def add_table_block(doc, table_lines):
    rows = []
    for line in table_lines:
        line = line.strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)
    if not rows:
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D6D0C4")
        tblBorders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    for row_idx, row_cells in enumerate(rows):
        for col_idx, cell_text in enumerate(row_cells):
            if col_idx >= n_cols:
                continue
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            cell_text = cell_text.replace("<br>", "\n").replace("\\|", "|")
            for j, segment in enumerate(cell_text.split("\n")):
                if j > 0:
                    p.add_run("\n")
                add_inline_runs(p, segment, base_size=10)
                if row_idx == 0:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = C_INK
            if row_idx == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F5F5F0")
                tcPr.append(shd)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_list_block(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        prefix = f"{i}. " if ordered else "• "
        r = p.add_run(prefix)
        set_run_font(r, size_pt=11, color=C_ORANGE, bold=True)
        add_inline_runs(p, item, base_size=11)


def render_block(doc, block, image_resolver, chapter_label=None,
                 is_first_h1=False):
    kind = block["kind"]
    if kind == "heading":
        level = block["level"]
        label = chapter_label if (level == 1 and is_first_h1) else None
        add_heading(doc, level, block["text"], chapter_label=label)
    elif kind == "p":
        if re.match(r"^\*[^*]+\*$", block["text"]):
            add_italic_subtitle(doc, block["text"].strip("*"))
        else:
            add_paragraph(doc, block["text"])
    elif kind == "code":
        add_code_block(doc, block["lines"])
    elif kind == "image":
        path = image_resolver(block)
        add_image(doc, path, block["alt"])
    elif kind == "quote":
        add_quote_block(doc, block["lines"])
    elif kind == "table":
        add_table_block(doc, block["rows"])
    elif kind == "ul":
        add_list_block(doc, block["items"], ordered=False)
    elif kind == "ol":
        add_list_block(doc, block["items"], ordered=True)
    elif kind == "hr":
        add_horizontal_line(doc)


# ============================================================
# 封面 / 目录
# ============================================================
def add_cover(doc, title, subtitle, author, extra_info=None):
    for _ in range(4):
        doc.add_paragraph()

    if extra_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(extra_info)
        set_run_font(r, size_pt=10, color=C_MUTED, font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size_pt=40, bold=True, color=C_INK,
                 font_cn=FONT_CN_HEAD, font_en="Georgia")
    p.paragraph_format.space_after = Pt(20)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run_font(r, size_pt=14, color=C_MUTED, font_cn=FONT_CN_HEAD)
        p.paragraph_format.space_after = Pt(60)

    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("作  者")
        set_run_font(r, size_pt=10, color=C_MUTED, font_cn=FONT_CN_HEAD)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(author)
        set_run_font(r, size_pt=18, bold=True, color=C_INK,
                     font_cn=FONT_CN_HEAD)

    add_page_break(doc)


def add_toc(doc, toc_items):
    """toc_items: list of (label, title)"""
    p = doc.add_paragraph()
    r = p.add_run("目  录")
    set_run_font(r, size_pt=20, bold=True, color=C_INK, font_cn=FONT_CN_HEAD)
    p.paragraph_format.space_after = Pt(20)

    for label, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.7
        p.paragraph_format.space_after = Pt(0)
        if label:
            r1 = p.add_run(label + "  ")
            set_run_font(r1, size_pt=10, color=C_MUTED, font_cn=FONT_CN_HEAD)
        r2 = p.add_run(title)
        set_run_font(r2, size_pt=12, color=C_INK)

    add_page_break(doc)


# ============================================================
# 页面设置
# ============================================================
def setup_page(doc, page_size="book"):
    section = doc.sections[0]
    if page_size == "book":
        # 大 32 开
        section.page_width = Cm(17.6)
        section.page_height = Cm(24.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    else:
        # A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = FONT_CN_BODY
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN_BODY)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)


def add_header_footer(doc, header_text=""):
    section = doc.sections[0]

    if header_text:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(header_text)
        set_run_font(r, size_pt=9, italic=True, color=C_MUTED,
                     font_cn=FONT_CN_HEAD)

    # 页脚页码
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    r._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r._r.append(fldChar2)
    set_run_font(r, size_pt=9, color=C_MUTED, font_cn=FONT_CN_HEAD)


# ============================================================
# 主流程
# ============================================================
def build_docx(md_files, output, images_dir=None, book_mode=False,
               title=None, subtitle=None, author=None, extra_info=None,
               chapter_labels=None, page_size="book"):
    doc = Document()
    setup_page(doc, page_size=page_size)

    if book_mode and title:
        add_header_footer(doc, header_text=title)
        add_cover(doc, title, subtitle, author, extra_info)

        # 目录：从每个 md 第一个 H1 提取
        toc_items = []
        for idx, md in enumerate(md_files):
            label = (chapter_labels[idx] if chapter_labels
                     and idx < len(chapter_labels)
                     else "")
            text = Path(md).read_text(encoding="utf-8")
            m = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
            chapter_title = m.group(1).strip() if m else Path(md).stem
            chapter_title = re.sub(r"^\d{2}\s+", "", chapter_title)
            toc_items.append((label, chapter_title))
        add_toc(doc, toc_items)

    elif book_mode:
        add_header_footer(doc, header_text="")

    for idx, md_path in enumerate(md_files):
        md_path = Path(md_path)
        if not md_path.exists():
            print(f"  ⚠ 跳过：{md_path}", file=sys.stderr)
            continue

        md_text = md_path.read_text(encoding="utf-8")
        image_refs = extract_image_refs(md_text)
        blocks = parse_blocks(md_text, image_refs)
        chapter_label = (chapter_labels[idx] if chapter_labels
                         and idx < len(chapter_labels) else None)

        # 图片解析器
        def make_resolver(md_p, refs):
            md_dir = md_p.parent

            def resolve(block):
                if block.get("via") == "inline":
                    p = block["path"]
                    cand = Path(p)
                    if not cand.is_absolute():
                        cand = md_dir / p
                    if not cand.exists() and images_dir:
                        cand = Path(images_dir) / Path(p).name
                    return str(cand) if cand.exists() else None
                else:
                    # 引用式：ref 映射到文件名或路径
                    ref = block.get("ref", "")
                    path_str = refs.get(ref)
                    if path_str:
                        cand = Path(path_str)
                        if not cand.is_absolute():
                            cand = md_dir / path_str
                        if not cand.exists() and images_dir:
                            cand = Path(images_dir) / Path(path_str).name
                        return str(cand) if cand.exists() else None
                    # 退化：直接按 ref 名字到 images_dir 找
                    if images_dir:
                        # 匹配 fig-1-1 → ch01-fig01.png
                        m = re.match(r"fig-(\d+)-(\d+)", ref)
                        if m:
                            ch_, fig_ = int(m.group(1)), int(m.group(2))
                            cand = Path(images_dir) / f"ch{ch_:02d}-fig{fig_:02d}.png"
                            if cand.exists():
                                return str(cand)
                    return None
            return resolve

        resolver = make_resolver(md_path, image_refs)

        n_imgs = sum(1 for b in blocks if b["kind"] == "image")
        text_chars = sum(len(b.get("text", "")) for b in blocks
                         if b["kind"] == "p")
        print(f"  ✓ {md_path.name}  ~{text_chars} 字 · {n_imgs} 图")

        is_first_h1 = True
        for b in blocks:
            render_block(doc, b, resolver, chapter_label=chapter_label,
                         is_first_h1=is_first_h1)
            if b["kind"] == "heading" and b["level"] == 1:
                is_first_h1 = False

        if book_mode and idx < len(md_files) - 1:
            add_page_break(doc)

    doc.save(output)
    size_mb = Path(output).stat().st_size / 1024 / 1024
    print(f"\n✨ 完成 · {output}  ({size_mb:.1f} MB)")


def main():
    ap = argparse.ArgumentParser(
        description="把 markdown 加工成出版社级 docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    ap.add_argument("md_files", nargs="+", help="一个或多个 md 文件")
    ap.add_argument("-o", "--output", help="输出 docx 路径（默认从首个 md 推导）")
    ap.add_argument("--images-dir", help="图片所在目录（默认从 md 同级目录找）")
    ap.add_argument("--book", action="store_true",
                    help="书籍模式：加封面+目录+页眉页脚+章节分页")
    ap.add_argument("--title", help="书名（书籍模式必填）")
    ap.add_argument("--subtitle", help="副标题")
    ap.add_argument("--author", help="作者")
    ap.add_argument("--extra-info", help="封面顶部小字（如「2026 年 · 橙皮书系列」）")
    ap.add_argument("--chapter-labels",
                    help="章号标签，逗号分隔（如「第 1 章,第 2 章,后记」）")
    ap.add_argument("--page-size", choices=["book", "a4"], default="book",
                    help="页面规格：book=大 32 开 / a4=A4")
    args = ap.parse_args()

    md_files = [Path(f) for f in args.md_files]
    for md in md_files:
        if not md.exists():
            print(f"❌ 找不到：{md}", file=sys.stderr)
            sys.exit(1)

    output = args.output
    if not output:
        if len(md_files) == 1:
            output = str(md_files[0].with_suffix(".docx"))
        else:
            output = "book.docx"

    chapter_labels = None
    if args.chapter_labels:
        chapter_labels = [s.strip() for s in args.chapter_labels.split(",")]

    if args.book and not args.title:
        print("❌ --book 模式必须指定 --title", file=sys.stderr)
        sys.exit(1)

    print(f"\n=== 构建 docx ===\n")
    build_docx(
        md_files=md_files,
        output=output,
        images_dir=args.images_dir,
        book_mode=args.book,
        title=args.title,
        subtitle=args.subtitle,
        author=args.author,
        extra_info=args.extra_info,
        chapter_labels=chapter_labels,
        page_size=args.page_size,
    )


if __name__ == "__main__":
    main()
